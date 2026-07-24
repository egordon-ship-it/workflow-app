"""
Extract latest email body from DM Marketing Word templates.

Version rule: if the doc contains headers like "V. 1" / "V. 2", keep only
the highest version block. Otherwise use the full document.
"""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document

ROOT = Path(r"v:\Marketing\Email Templates")
OUT_DIR = Path(
    r"c:\Users\egordon\Desktop\Cursor Projects\workflow builder\lib\emails\content"
)
OUT_JSON = Path(
    r"c:\Users\egordon\Desktop\Cursor Projects\workflow builder\scripts\extracted-emails.json"
)

# Explicit path list from the user (relative to ROOT)
REL_PATHS = [
    r"Customers\CCDenRem3.docx",
    r"Customers\CCExpire1.docx",
    r"Customers\CCExpire2.docx",
    r"Customers\DMSTCMSYB.docx",
    r"Customers\DMSTCONFRM.docx",
    r"Customers\InvoiceReq.docx",
    r"Customers\InvoiceRet.docx",
    r"Customers\InvRem1.docx",
    r"Customers\InvRem2.docx",
    r"Customers\MsgDemoWlc.docx",
    r"Customers\OrderShip.docx",
    r"Customers\PayInvoice.docx",
    r"Customers\PayInvText - SMS.docx",
    r"Customers\PortPassRes.docx",
    r"Customers\PortWelcome.docx",
    r"Customers\PrtDlrPR.docx",
    r"Customers\PrtDlrWlc.docx",
    r"Customers\PSAVFDBKCONF.docx",
    r"Customers\PSAVSUBWARN.docx",
    r"Customers\REACTDEMO.docx",
    r"Customers\RecACHRej.docx",
    r"Customers\RecCCApp.docx",
    r"Customers\REJECTDEMO.docx",
    r"Customers\SATSENDCRED.docx",
    r"Customers\SENDCRED.docx",
    r"Customers\SENDREBATE.docx",
    r"Customers\SENDWRNTY.docx",
    r"Customers\SndDlrOnSgn.docx",
    r"Customers\SndOnSgn.docx",
    r"Customers\SNDPARTLINKS.docx",
    r"Customers\SndPrpsl.docx",
    r"Customers\SURVEYINV.docx",
    r"Customers\SYBACT.docx",
    r"Customers\SYBDACTRYL.docx",
    r"Customers\SYBFREETRYL.docx",
    r"Customers\UpdateCC.docx",
    r"Customers\ACTDEMO.docx",
    r"Customers\CCDenRem1.docx",
    r"Customers\CCDenRem2.docx",
    r"Encore\PSAVNEWSITE.docx",
    r"Encore\PSAVNEWUSER.docx",
    r"Encore\PSAVSUBENDED.docx",
    r"Encore\PSAVSUBWARN.docx",
    r"Encore\PSAVSVCACTN.docx",
    r"Encore\PSAVUSERFAIL.docx",
    r"Encore\PSAVUSERREM1.docx",
    r"Encore\PSAVUSERREM2.docx",
    r"Encore\PSAVUSERREM3.docx",
    r"Internal\AEPayNotify.docx",
    r"Internal\BIGCUSTTASK.docx",
    r"Internal\CCREFUNDERR.docx",
    r"Internal\DEACTDMX.docx",
    r"Internal\DEACTPLAY.docx",
    r"Internal\DLRSGNCMPLT.docx",
    r"Internal\DMPSSFUM.docx",
    r"Internal\LDINSRTASGN.docx",
    r"Internal\NEWJOBENTRY.docx",
    r"Internal\NOTFYNEWRENO - TH.docx",
    r"Internal\NOTIFYDEMO.docx",
    r"Internal\PSAVFDBKSBMT - Encore.docx",
    r"Internal\SPNOTIFY - TH.docx",
    r"Internal\SURVEYCMPLT.docx",
    r"Internal\SXMSVCERROR.docx",
    r"Internal\SYBCTNOTIFY.docx",
    r"Internal\SYBFT2PAID.docx",
    r"Internal\SYBNOTDEMO.docx",
    r"Internal\SYBNOTIFY.docx",
    r"Internal\SYBSVCERROR.docx",
    r"Internal\TASKASSIGN.docx",
    r"Internal\TASKREASSIGN.docx",
    r"Internal\TSKCONFLICT.docx",
    r"Internal\XMSTORABAND.docx",
    r"TechniciansDealers\ACTSXM.docx",
    r"TechniciansDealers\BXSWPSXM.docx",
    r"TechniciansDealers\DEACTSXIR.docx",
    r"TechniciansDealers\DEACTSXM.docx",
    r"TechniciansDealers\DEALERINV.docx",
    r"TechniciansDealers\DealerUpdat.docx",
    r"TechniciansDealers\FREETRIAL.docx",
    r"TechniciansDealers\PrtDlrWlcCM.docx",
    r"TechniciansDealers\PrtDlrWlcDB.docx",
    r"TechniciansDealers\REACTSXIR.docx",
    r"TechniciansDealers\REACTSXM.docx",
    r"TechniciansDealers\RESENDTRIAL.docx",
    r"TechniciansDealers\ACTDMX.docx",
    r"TechniciansDealers\ACTPLAY.docx",
    r"TechniciansDealers\ACTSXIR.docx",
    r"Tim Hortons\SHOPTHCONF.docx",
    r"Tim Hortons\SHOPTHCONFF.docx",
    r"Tim Hortons\THGETSTART.docx",
    r"Tim Hortons\THGETSTARTFR.docx",
    r"Tim Hortons\THNOTONLINE.docx",
    r"Tim Hortons\THNOTONLINEFR.docx",
    r"Tim Hortons\THOrderShip.docx",
    r"Tim Hortons\THPLONLINE.docx",
    r"Tim Hortons\THPLONLINEFR.docx",
    r"Tim Hortons\THREMOTECODE.docx",
    r"Tim Hortons\PassResTH.docx",
    r"Tim Hortons\PassResTHFR.docx",
    r"Tim Hortons\PRTLTHWLCM.docx",
    r"Tim Hortons\PRTLTHWLCMFR.docx",
]

# Prefer these folders when the same CRM key appears twice
PREFER_FOLDER = {
    "PSAVSUBWARN": "Encore",
}

VERSION_RE = re.compile(
    r"^\s*V\.?\s*(\d+)\s*$",
    re.IGNORECASE,
)
# Also match inline headers like "V. 2 — Subject" or "Version 2"
VERSION_LOOSE_RE = re.compile(
    r"^\s*(?:V(?:ersion)?\.?\s*|Ver\.?\s*)(\d+)\b",
    re.IGNORECASE,
)


def filename_to_key(name: str) -> str:
    stem = Path(name).stem
    # Strip suffixes like " - SMS", " - TH", " - Encore"
    stem = re.sub(r"\s+-\s+.+$", "", stem)
    return stem.strip()


def paragraph_texts(doc: Document) -> list[str]:
    lines: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").replace("\xa0", " ").rstrip()
        # Keep blank lines as paragraph breaks (collapse later)
        lines.append(t)
    # Also pull simple tables (common in these logs)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").replace("\xa0", " ").strip() for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                lines.append(" | ".join(cells))
    return lines


def split_versions(lines: list[str]) -> tuple[int | None, list[str]]:
    """
    Returns (version_number_or_None, body_lines).
    If multiple V.n headers exist, keep the highest n's body until next V.* or EOF.
    """
    markers: list[tuple[int, int]] = []  # (version, line_index)
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
        m = VERSION_RE.match(stripped) or VERSION_LOOSE_RE.match(stripped)
        if m:
            # Avoid treating "V. 2 of something long body text" mid-paragraph as a
            # marker unless the line is short (header-like)
            if len(stripped) <= 40:
                markers.append((int(m.group(1)), i))

    if not markers:
        return None, lines

    # Dedupe: keep last index per version number if repeated
    by_ver: dict[int, int] = {}
    for ver, idx in markers:
        by_ver[ver] = idx

    latest = max(by_ver.keys())
    start = by_ver[latest] + 1  # skip the version header itself

    # End at the next higher-or-any other version marker after start
    end = len(lines)
    for ver, idx in markers:
        if idx > by_ver[latest]:
            end = min(end, idx)

    return latest, lines[start:end]


def normalize_body(lines: list[str]) -> str:
    # Trim leading/trailing blanks; collapse 3+ blanks to 2
    text = "\n".join(lines)
    text = text.strip("\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_subject_hint(body: str) -> str | None:
    for line in body.splitlines()[:12]:
        s = line.strip()
        if not s:
            continue
        if re.match(r"^subject\s*:", s, re.I):
            return re.sub(r"^subject\s*:\s*", "", s, flags=re.I).strip()
    return None


def main() -> None:
    results: dict[str, dict] = {}
    errors: list[str] = []

    for rel in REL_PATHS:
        path = ROOT / rel
        folder = path.parent.name
        key = filename_to_key(path.name)
        if not path.exists():
            errors.append(f"MISSING {path}")
            continue

        # Preference when duplicate keys
        if key in results:
            prefer = PREFER_FOLDER.get(key)
            if prefer and folder != prefer:
                continue
            if prefer and results[key].get("folder") == prefer:
                continue

        try:
            doc = Document(str(path))
            lines = paragraph_texts(doc)
            version, body_lines = split_versions(lines)
            body = normalize_body(body_lines)
            subject_hint = extract_subject_hint(body)
            results[key] = {
                "key": key,
                "sourcePath": str(path),
                "folder": folder,
                "versionUsed": version,
                "subjectHint": subject_hint,
                "body": body,
                "charCount": len(body),
                "lineCount": len(body.splitlines()) if body else 0,
            }
        except Exception as e:
            errors.append(f"ERROR {path}: {e}")

    OUT_JSON.parent.mkdir(parents=True, exist_ok=True)
    OUT_JSON.write_text(
        json.dumps({"emails": results, "errors": errors, "count": len(results)}, indent=2),
        encoding="utf-8",
    )

    # Per-key JSON files for the app
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for key, data in results.items():
        (OUT_DIR / f"{key}.json").write_text(
            json.dumps(
                {
                    "key": data["key"],
                    "versionUsed": data["versionUsed"],
                    "sourcePath": data["sourcePath"],
                    "subjectHint": data["subjectHint"],
                    "body": data["body"],
                },
                indent=2,
                ensure_ascii=False,
            )
            + "\n",
            encoding="utf-8",
        )

    print(f"Extracted {len(results)} emails → {OUT_DIR}")
    print(f"Manifest → {OUT_JSON}")
    if errors:
        print(f"{len(errors)} errors:")
        for e in errors:
            print(" ", e)

    # Coverage summary
    with_v = sum(1 for d in results.values() if d["versionUsed"] is not None)
    empty = [k for k, d in results.items() if not d["body"].strip()]
    print(f"With version header: {with_v}")
    print(f"Empty bodies: {len(empty)} {empty[:10]}")


if __name__ == "__main__":
    main()
